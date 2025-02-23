# import os
# import json
# import boto3
# import pdfplumber
# import docx
# from fpdf import FPDF
# from flask import Flask, render_template, request, send_file
# from werkzeug.utils import secure_filename

# UPLOAD_FOLDER = "uploads"
# RESULTS_FOLDER = "results"
# ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}
# AWS_REGION = "ap-south-1"
# BEDROCK_MODEL_ID = "meta.llama3-8b-instruct-v1:0"

# app = Flask(__name__)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# app.config["RESULTS_FOLDER"] = RESULTS_FOLDER
# app.config["ALLOWED_EXTENSIONS"] = ALLOWED_EXTENSIONS

# bedrock = boto3.client("bedrock-runtime", region_name=AWS_REGION)

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config["ALLOWED_EXTENSIONS"]

# def extract_text_from_file(file_path):
#     text = ""
#     if file_path.endswith(".pdf"):
#         with pdfplumber.open(file_path) as pdf:
#             text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
#     elif file_path.endswith(".docx"):
#         doc = docx.Document(file_path)
#         text = "\n".join([para.text for para in doc.paragraphs])
#     elif file_path.endswith(".txt"):
#         with open(file_path, "r", encoding="utf-8") as file:
#             text = file.read()
#     return text.strip()

# def generate_mcqs(text, num_mcqs, difficulty, num_options):
#     prompt = (f"Generate {num_mcqs} multiple-choice questions at {difficulty} difficulty level "
#               f"with {num_options} answer options from the following text:\n\n{text}\n\n")
#     payload = {"prompt": prompt, "max_gen_len": 1024, "temperature": 0.5, "top_p": 0.9}
#     response = bedrock.invoke_model(
#         modelId=BEDROCK_MODEL_ID,
#         body=json.dumps(payload),
#         contentType="application/json",
#         accept="application/json"
#     )
#     response_body = json.loads(response["body"].read())
#     mcqs = response_body.get("generation", "No MCQs generated.").strip().split("\n")  # Convert to list
#     return mcqs

# def create_pdf(mcqs, filename):
#     pdf = FPDF()
#     pdf.set_auto_page_break(auto=True, margin=15)
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     for line in mcqs:
#         pdf.multi_cell(0, 10, line)
#     pdf_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
#     pdf.output(pdf_path)
#     return pdf_path

# @app.route("/")
# def home():
#     return render_template("index.html")

# @app.route("/upload", methods=["POST"])
# def upload():
#     if "file" not in request.files:
#         return "No file uploaded"
#     file = request.files["file"]
#     if file and allowed_file(file.filename):
#         filename = secure_filename(file.filename)
#         file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
#         file.save(file_path)
#         text = extract_text_from_file(file_path)
#         num_mcqs = int(request.form.get("num_mcqs", 5))
#         difficulty = request.form.get("difficulty", "Medium")
#         num_options = int(request.form.get("num_options", 4))
#         mcqs_list = generate_mcqs(text, num_mcqs, difficulty, num_options)
#         pdf_filename = f"{filename}_mcqs.pdf"
#         pdf_path = create_pdf(mcqs_list, pdf_filename)
#         return render_template("results.html", mcqs=mcqs_list, pdf_filename=pdf_filename)
#     return "Invalid file format"

# @app.route("/download/<filename>")
# def download(filename):
#     file_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
#     return send_file(file_path, as_attachment=True)

# if __name__ == "__main__":
#     os.makedirs(UPLOAD_FOLDER, exist_ok=True)
#     os.makedirs(RESULTS_FOLDER, exist_ok=True)
#     app.run(debug=True, port=6001)
import os
import json
import boto3
import pdfplumber
import docx
from fpdf import FPDF
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename

UPLOAD_FOLDER = "uploads"
RESULTS_FOLDER = "results"
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}
AWS_REGION = "ap-south-1"
BEDROCK_MODEL_ID = "meta.llama3-8b-instruct-v1:0"

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["RESULTS_FOLDER"] = RESULTS_FOLDER
app.config["ALLOWED_EXTENSIONS"] = ALLOWED_EXTENSIONS

bedrock = boto3.client("bedrock-runtime", region_name=AWS_REGION)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config["ALLOWED_EXTENSIONS"]

def extract_text_from_file(file_path):
    text = ""
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
    return text.strip()

def generate_mcqs(text, num_mcqs, difficulty, num_options):
    prompt = (
        f"Generate {num_mcqs} multiple-choice questions at {difficulty} difficulty level "
        f"with {num_options} answer options from the following text:\n\n{text}\n\n"
        "Provide only structured MCQs in this format:\n"
        "Q1: [Question text]\nA) Option 1\nB) Option 2\nC) Option 3\nD) Option 4\nAnswer: [Correct Option]"
    )
    
    payload = {"prompt": prompt, "max_gen_len": 1024, "temperature": 0.5, "top_p": 0.9}
    response = bedrock.invoke_model(
        modelId=BEDROCK_MODEL_ID,
        body=json.dumps(payload),
        contentType="application/json",
        accept="application/json"
    )
    
    response_body = json.loads(response["body"].read())
    raw_mcqs = response_body.get("generation", "").strip()
    
    # Ensure proper formatting by extracting only valid MCQs
    mcqs = []
    lines = raw_mcqs.split("\n")
    for line in lines:
        line = line.strip()
        if line and not line.startswith("**"):  # Remove unwanted formatting
            mcqs.append(line)
    
    return mcqs

def create_pdf(mcqs, filename):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for line in mcqs:
        pdf.multi_cell(0, 10, line)
    
    pdf_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
    pdf.output(pdf_path)
    return pdf_path

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return "No file uploaded"
    
    file = request.files["file"]
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(file_path)
        
        text = extract_text_from_file(file_path)
        num_mcqs = int(request.form.get("num_mcqs", 5))
        difficulty = request.form.get("difficulty", "Medium")
        num_options = int(request.form.get("num_options", 4))
        
        mcqs_list = generate_mcqs(text, num_mcqs, difficulty, num_options)
        pdf_filename = f"{filename}_mcqs.pdf"
        pdf_path = create_pdf(mcqs_list, pdf_filename)
        
        return render_template("results.html", mcqs=mcqs_list, pdf_filename=pdf_filename)
    
    return "Invalid file format"

@app.route("/download/<filename>")
def download(filename):
    file_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
    return send_file(file_path, as_attachment=True)

if __name__ == "__main__":
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(RESULTS_FOLDER, exist_ok=True)
    app.run(debug=True, port=6001)
